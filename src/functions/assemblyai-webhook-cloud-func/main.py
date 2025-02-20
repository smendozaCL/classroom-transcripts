import json
import logging
import os
from flask import jsonify
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
import assemblyai as aai
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from google.cloud import storage
import google.auth
from datetime import datetime

load_dotenv()

# Initialize AssemblyAI client if API key is available
api_key = os.environ.get("ASSEMBLYAI_API_KEY")
if api_key:
    aai.settings.api_key = api_key
    logging.info("AssemblyAI API key configured")
else:
    logging.info("No AssemblyAI API key found - speaker diarization will be disabled")


def verify_webhook_signature(request):
    """Verify the webhook signature from AssemblyAI"""
    webhook_auth_value = os.environ.get("ASSEMBLYAI_WEBHOOK_AUTH_HEADER_VALUE")
    if not webhook_auth_value:
        logging.error(
            "ASSEMBLYAI_WEBHOOK_AUTH_HEADER_VALUE environment variable not set"
        )
        return False

    # Log all headers for debugging
    logging.info("Request headers:")
    for header_name, header_value in request.headers.items():
        logging.info(f"  {header_name}: {header_value}")

    # Verify the request is coming from AssemblyAI's IP
    forwarded_for = request.headers.get("X-Forwarded-For", "")
    if "44.238.19.20" not in forwarded_for:
        logging.error(f"Request not from AssemblyAI IP. Got: {forwarded_for}")
        return False

    # Verify User-Agent
    user_agent = request.headers.get("User-Agent", "")
    if user_agent != "AssemblyAI-Webhook":
        logging.error(f"Invalid User-Agent. Got: {user_agent}")
        return False

    # Get the auth header value - case insensitive search
    received_auth = None
    for header_name, header_value in request.headers.items():
        if (
            header_name.lower() == "x-transcript-webhook-secret".lower()
        ):  # Match what we configured
            received_auth = header_value
            break

    if not received_auth:
        logging.error("No webhook auth header found")
        logging.error("Headers received:")
        for name, value in request.headers.items():
            logging.error(f"  {name}: {value}")
        return False

    # Simple string comparison of the auth values
    if received_auth != webhook_auth_value:
        logging.error(
            f"Invalid webhook auth. Expected length: {len(webhook_auth_value)}, Got length: {len(received_auth)}"
        )
        return False

    return True


def upload_to_drive(transcript_data, transcript_id):
    """Upload transcript to Google Drive"""
    try:
        # Initialize Drive API with default credentials
        SCOPES = ["https://www.googleapis.com/auth/drive.file"]
        credentials, project = google.auth.default(scopes=SCOPES)
        drive_service = build(
            "drive", "v3", credentials=credentials, cache_discovery=False
        )
        folder_id = os.environ["DRIVE_FOLDER_ID"]
        logging.info(f"Using Drive folder ID: {folder_id}")

        # Generate and upload Word document
        docx_metadata = {
            "name": f"{transcript_data.get('raw_webhook_data', {}).get('original_filename', 'transcript')} - {datetime.now().strftime('%Y-%m-%d %H:%M')} - {'With Speaker Labels' if 'utterances' in transcript_data else 'No Speaker Labels'}.docx",
            "parents": [folder_id],
        }
        docx_content = BytesIO(generate_transcript_docx(transcript_data))
        docx_media = MediaIoBaseUpload(
            docx_content,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            resumable=True,
        )
        docx_file = (
            drive_service.files()
            .create(body=docx_metadata, media_body=docx_media, fields="id")
            .execute()
        )

        logging.info(f"Word document uploaded to Drive with ID: {docx_file.get('id')}")
        return docx_file.get("id")

    except Exception as e:
        logging.error(f"Failed to upload to Drive: {str(e)}")
        raise


def generate_transcript_docx(transcript_data):
    """
    Generate a docx file from transcript data.

    Args:
        transcript_data: Dictionary containing transcript information

    Returns:
        bytes: The generated docx file as bytes
    """
    doc = Document()
    doc.add_heading("Transcript", 0)

    # Get transcript from AssemblyAI to check for speaker diarization
    try:
        transcript = aai.Transcript.get_by_id(transcript_data["transcript_id"])

        if transcript.utterances:
            # Add each speaker's text as a paragraph
            for utterance in transcript.utterances:
                p = doc.add_paragraph()
                # Add timestamp
                start_seconds = (
                    utterance.start / 1000.0
                )  # Convert milliseconds to seconds
                hours = int(start_seconds // 3600)
                minutes = int((start_seconds % 3600) // 60)
                seconds = int(start_seconds % 60)
                timestamp = f"[{hours:02d}:{minutes:02d}:{seconds:02d}]"

                # Add speaker label in bold
                speaker_letter = (
                    chr(65 + (utterance.speaker - 1))
                    if isinstance(utterance.speaker, int)
                    else utterance.speaker
                )
                speaker_run = p.add_run(f"{timestamp} Speaker {speaker_letter}: ")
                speaker_run.bold = True

                # Add the text
                p.add_run(utterance.text)
                # Add spacing between utterances
                p.add_run("\n")
        else:
            # Add the full text as a single paragraph
            doc.add_paragraph(transcript_data["text"])

    except Exception as e:
        logging.warning(
            f"Could not get speaker diarization, using plain text: {str(e)}"
        )
        # If we can't get the transcript from AssemblyAI, just use the text we have
        doc.add_paragraph(transcript_data["text"])

    # Add metadata
    doc.core_properties.title = "Transcript"
    doc.core_properties.comments = "Generated from AssemblyAI transcription"

    # Save to bytes
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()


def handle_assemblyai_webhook(request):
    """
    Google Cloud Function to handle AssemblyAI webhooks
    Args:
        request (flask.Request): The request object
    Returns:
        flask.Response: The response object
    """
    logging.info("AssemblyAI Webhook triggered")

    # Get raw request data first
    raw_data = request.get_data()
    if not raw_data:
        logging.error("No request data received")
        return jsonify({"error": "No request data"}), 400

    logging.info(f"Raw request data: {raw_data.decode('utf-8')}")

    # Parse JSON data
    try:
        webhook_data = json.loads(raw_data)
        logging.info(f"Parsed webhook data: {json.dumps(webhook_data)}")
    except json.JSONDecodeError as e:
        logging.error(f"Failed to parse webhook data: {str(e)}")
        return jsonify({"error": "Invalid JSON data"}), 400

    # Verify this is a transcript_completed event
    if webhook_data.get("status") != "completed":
        return jsonify({"message": "Ignoring non-completed transcript status"}), 200

    # Get transcript ID
    transcript_id = webhook_data.get("transcript_id")
    if not transcript_id:
        logging.error(f"Missing transcript_id in webhook data: {webhook_data}")
        return jsonify({"error": "Missing transcript_id"}), 400

    try:
        # Get the full transcript from AssemblyAI
        transcript = aai.Transcript.get_by_id(transcript_id)
        if not transcript or not transcript.text:
            logging.error(f"Could not retrieve transcript {transcript_id}")
            return jsonify({"error": "Could not retrieve transcript"}), 500

        # Initialize Google Cloud Storage client
        if not os.environ.get("GOOGLE_CLOUD_PROJECT"):
            raise ValueError("GOOGLE_CLOUD_PROJECT environment variable is not set")

        storage_client = storage.Client()
        bucket = storage_client.bucket(os.environ["BUCKET_NAME"])

        # Prepare transcript data
        transcript_data = {
            "transcript_id": transcript_id,
            "text": transcript.text,
            "status": "completed",
            "raw_webhook_data": webhook_data,
        }

        # Store JSON transcript
        blob_name = f"transcripts/{transcript_id}/transcript.json"
        blob = bucket.blob(blob_name)
        blob.upload_from_string(
            json.dumps(transcript_data), content_type="application/json"
        )
        logging.info(f"Successfully stored JSON transcript in GCS: {transcript_id}")

        # Store Word document
        docx_blob_name = f"transcripts/{transcript_id}/transcript.docx"
        docx_blob = bucket.blob(docx_blob_name)
        docx_blob.upload_from_string(
            generate_transcript_docx(transcript_data),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        logging.info(f"Successfully stored Word document in GCS: {transcript_id}")

        # Also store Word document in Drive
        drive_file_id = upload_to_drive(transcript_data, transcript_id)
        logging.info(f"Successfully stored Word document in Drive: {drive_file_id}")

        return jsonify(
            {
                "message": "Transcript processed successfully",
                "gcs_paths": {"json": blob_name, "docx": docx_blob_name},
                "drive_file_id": drive_file_id,
            }
        ), 200

    except Exception as e:
        logging.error(f"Error processing webhook: {str(e)}")
        return jsonify({"error": str(e)}), 500
