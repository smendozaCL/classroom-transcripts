import streamlit as st
from src.utils.view_table import get_table_client, list_table_items
from datetime import datetime
import pytz
import assemblyai as aai
import os
import logging
from docx import Document
from io import BytesIO
import pydantic

from utils.azure_storage import get_sas_url_for_audio_file_name

DEBUG = bool(st.secrets.get("DEBUG", False))
table_name = st.session_state.get(
    "table_name", st.secrets.get("AZURE_STORAGE_TABLE_NAME")
)
TRANSCRIPT_PREVIEW_MAX_LENGTH = 1000
TRANSCRIPT_PREVIEW_SPEAKER_TURNS = 5

if not st.experimental_user.get("is_logged_in"):
    st.login()

# Initialize session state for status values if not already set
if "transcription_statuses" not in st.session_state:
    st.session_state.transcription_statuses = [
        "queued",  # Initial state
        "processing",  # Being transcribed
        "completed",  # Done
        "error",  # Failed
        "failed",  # Another error state
    ]

# US timezone options
US_TIMEZONES = [
    "US/Eastern",
    "US/Central",
    "US/Mountain",
    "US/Pacific",
    "US/Alaska",
    "US/Hawaii",
]

# Initialize timezone in session state
if "timezone" not in st.session_state:
    st.session_state.timezone = "US/Pacific"

# Add timezone picker to sidebar
with st.sidebar:
    st.session_state.timezone = st.selectbox(
        "Select Timezone",
        options=US_TIMEZONES,
        index=US_TIMEZONES.index(st.session_state.timezone),
        help="Choose your local timezone",
        format_func=lambda x: x.replace("US/", ""),
    )

# Update local_tz to use session state
local_tz = pytz.timezone(st.session_state.timezone)


# Initialize AssemblyAI client
aai.settings.api_key = os.getenv("ASSEMBLYAI_API_KEY")
transcriber = aai.Transcriber()

# Initialize session state for pagination
if "items_per_page" not in st.session_state:
    st.session_state.items_per_page = 5  # Initial number of items to show
if "current_page" not in st.session_state:
    st.session_state.current_page = 1

# Get admin emails from Streamlit secrets
ADMIN_EMAILS = [
    email.strip().lower() for email in st.secrets.get("admin_emails", "").split(",")
]

# Debug logging for admin list
if DEBUG:
    st.write("Debug - Admin emails:", ADMIN_EMAILS)


def is_admin(email: str) -> bool:
    """Check if the given email belongs to an admin"""
    return email.lower() in ADMIN_EMAILS


def reset_pagination():
    """Reset pagination state"""
    st.session_state.items_per_page = 5
    st.session_state.current_page = 1


def format_file_size(size_in_bytes):
    """Convert bytes to human readable format"""
    if not isinstance(size_in_bytes, (int, float)):
        return "N/A"
    for unit in ["B", "KB", "MB", "GB"]:
        if size_in_bytes < 1024:
            return f"{size_in_bytes:.2f} {unit}"
        size_in_bytes /= 1024
    return f"{size_in_bytes:.2f} TB"


# Get timezone abbreviation


def localized_timestamp(timestamp):
    """Get localized timestamp"""
    local_timestamp = timestamp.astimezone(local_tz)
    time_string = local_timestamp.strftime("%Y-%m-%d %H:%M:%S %Z")
    return time_string


def get_transcript_status(transcript_id):
    """Get transcript status from AssemblyAI"""
    # Skip test data
    if transcript_id.startswith("test_"):
        return "completed"

    try:
        transcript = aai.Transcript.get_by_id(transcript_id)
        return transcript.status.value
    except Exception as e:
        if "not found" in str(e).lower():
            return "error"  # Transcript doesn't exist in AssemblyAI
        st.error(f"Error getting transcript status: {str(e)}")
        return "error"


def generate_transcript_markdown(transcript, max_length=None, max_speaker_turns=None):
    """
    Generate markdown formatted text from an AssemblyAI transcript.

    Args:
        transcript: AssemblyAI Transcript object
        max_length (int, optional): Maximum length of text to include
        max_speaker_turns (int, optional): Maximum number of speaker turns to include

    Returns:
        str: Markdown formatted transcript text
    """
    if not transcript:
        return "No transcript available"

    markdown_lines = []

    # Handle transcripts with speaker detection
    if transcript.utterances:
        for i, utterance in enumerate(transcript.utterances):
            # Break if we've hit the max speaker turns
            if max_speaker_turns and i >= max_speaker_turns:
                markdown_lines.append(
                    "\n*[Additional transcript content truncated...]*"
                )
                break

            # Format timestamp as [00:00:00]
            start_seconds = utterance.start / 1000.0  # Convert milliseconds to seconds
            hours = int(start_seconds // 3600)
            minutes = int((start_seconds % 3600) // 60)
            seconds = int(start_seconds % 60)
            timestamp = f"[{hours:02d}:{minutes:02d}:{seconds:02d}]"

            # Format as [timestamp] **Speaker X**: text
            speaker_letter = (
                chr(65 + (utterance.speaker - 1))
                if isinstance(utterance.speaker, int)
                else utterance.speaker
            )
            speaker_text = f"{timestamp} **Speaker {speaker_letter}**: {utterance.text}"
            markdown_lines.append(speaker_text)

            # Check total length if max_length specified
            current_text = "\n".join(markdown_lines)
            if max_length and len(current_text) >= max_length:
                truncate_length = max_length - len(
                    "\n\n*[Additional transcript content truncated...]*"
                )
                markdown_lines[-1] = str(markdown_lines[-1])[:truncate_length]
                markdown_lines.append(
                    "\n*[Additional transcript content truncated...]*"
                )
                break

    # Handle transcripts without speaker detection
    elif transcript.text:
        text = transcript.text
        if max_length:
            truncate_length = max_length
            text = str(text)[:truncate_length] + (
                "..." if len(transcript.text) > truncate_length else ""
            )
        markdown_lines.append(text)

    return "\n\n".join(markdown_lines)


def generate_transcript_docx(transcript):
    """
    Generate a docx file from an AssemblyAI transcript.

    Args:
        transcript: AssemblyAI Transcript object

    Returns:
        bytes: The generated docx file as bytes
    """
    doc = Document()
    doc.add_heading("Transcript", 0)

    if transcript.utterances:
        # Add each speaker's text as a paragraph
        for utterance in transcript.utterances:
            p = doc.add_paragraph()
            # Add speaker label in bold
            speaker_run = p.add_run(f"Speaker {utterance.speaker}: ")
            speaker_run.bold = True
            # Add the text
            p.add_run(utterance.text)
            # Add spacing between utterances
            p.add_run("\n")
    else:
        # Add the full text as a single paragraph
        doc.add_paragraph(transcript.text)

    # Add metadata
    doc.core_properties.title = "Transcript"
    doc.core_properties.comments = "Generated from AssemblyAI transcription"

    # Save to bytes
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()


st.title("🔍 Audio Files & Transcriptions")

# Initialize table client
table_client = get_table_client(table_name)

if DEBUG:
    st.write(
        f"Debug - Table name: {st.session_state.get('table_name', 'Transcriptions')}"
    )


def can_view_transcript(transcript_email: str, user_email: str) -> bool:
    """Check if user can view a specific transcript"""
    if is_admin(user_email):
        return True
    # If no uploader email is set, only admins can view
    if not transcript_email:
        return False
    return transcript_email.lower() == user_email.lower()


def get_transcript_statuses():
    """Get all transcript statuses in one API call"""
    # Check if we have cached statuses and they're not expired

    try:
        status_map = {}
        params = aai.ListTranscriptParameters(limit=100)  # Adjust limit as needed

        # Get first page
        page = transcriber.list_transcripts(params)
        for t in page.transcripts:
            # Handle test data
            if t.id.startswith("test_"):
                status_map[t.id] = "completed"
            else:
                status_map[t.id] = t.status.value

        # Paginate through all remaining pages
        while page.page_details.before_id_of_prev_url is not None:
            params.before_id = page.page_details.before_id_of_prev_url
            page = transcriber.list_transcripts(params)
            for t in page.transcripts:
                if t.id.startswith("test_"):
                    status_map[t.id] = "completed"
                else:
                    status_map[t.id] = t.status.value

        return status_map
    except Exception as e:
        st.error(f"Error getting transcript statuses: {str(e)}")
        # Return cached data if available, even if expired
        return {}


def query_table_entities(table_client, user_email: str):
    """
    Query table entities based on user permissions.

    Args:
        table_client: Azure TableClient instance
        user_email: Email of the current user

    Returns:
        List of entities the user has permission to view
    """

    if not user_email:
        return []

    try:
        # For regular users, only fetch their items
        if not is_admin(user_email):
            if DEBUG:
                st.info(
                    f"Debug - User {user_email} is not admin, fetching only their items"
                )
            filter_condition = f"uploaderEmail eq '{user_email.lower()}'"
            items = list(table_client.query_entities(filter_condition))
        else:
            if DEBUG:
                st.info(f"Debug - User {user_email} is admin, fetching all items")
            items = list_table_items(
                st.session_state.get(
                    "table_name", st.secrets.get("AZURE_STORAGE_TABLE_NAME")
                )
            )

        if DEBUG:
            st.info(f"Debug - Number of items fetched: {len(items) if items else 0}")
        return items

    except Exception as e:
        logging.error(f"Error querying table: {e}")
        if DEBUG:
            st.error(f"Debug - Error querying table: {str(e)}")
            st.write(f"Debug - Table client state: {table_client}")
        return []


def load_table_data(_table_client):
    """Load and process table data with caching"""
    MIN_DATE = datetime(2000, 1, 1, tzinfo=pytz.UTC)

    user = st.experimental_user
    validated_email = user.email if user.email_verified else None

    if validated_email is not None:
        # Use consolidated query function
        items = query_table_entities(_table_client, str(validated_email))

    if not items:
        return []

    # Get all transcript statuses at once
    transcript_statuses = get_transcript_statuses()
    items_list = []

    for item in items:
        item_dict = dict(item)

        # Add formatted size
        if "blobSize" in item_dict:
            item_dict["formatted_size"] = format_file_size(item_dict["blobSize"])

        # Get status from cached transcript statuses
        if "transcriptId" in item_dict:
            item_dict["status"] = transcript_statuses.get(
                item_dict["transcriptId"], "error"
            )
        else:
            item_dict["status"] = "pending"

        item_dict["_previous_status"] = item_dict["status"]

        # Process timestamp
        if "uploadTime" not in item_dict:
            item_dict["uploadTime"] = item_dict.get("Timestamp", MIN_DATE)

        try:
            # Handle different timestamp types
            if isinstance(item_dict["uploadTime"], str):
                dt = datetime.fromisoformat(
                    item_dict["uploadTime"].replace("Z", "+00:00")
                )
            elif isinstance(item_dict["uploadTime"], datetime):
                dt = item_dict["uploadTime"]
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=pytz.UTC)
            else:
                dt = MIN_DATE

            local_dt = dt.astimezone(local_tz)
            item_dict["_timestamp"] = local_dt
            item_dict["uploadTime"] = local_dt
        except ValueError as e:
            logging.error(f"Error parsing time: {e}")
            item_dict["_timestamp"] = MIN_DATE
            item_dict["uploadTime"] = MIN_DATE

        # Add class name and description
        item_dict["className"] = item_dict.get("className", None)
        item_dict["description"] = item_dict.get("description", None)

        items_list.append(item_dict)

    return items_list


def display_transcript_item(item):
    """Display a single transcript item in a fragment"""
    try:
        transcript_id = item.get("transcriptId")
        if not transcript_id:
            st.error("Missing transcript ID")
            return

        # Get status info for formatting
        upload_time = item.get("uploadTime")
        upload_time_str = localized_timestamp(upload_time)
        original_file_name = item.get("originalFileName", "Untitled")
        row_key = item.get("RowKey", "")
        status = item.get("status")
        class_name = item.get("className", "")
        uploader_email = item.get("uploaderEmail", "Unknown")

        # Choose icon based on status
        status_icon = "📄"  # Default icon
        if status in ["queued", "processing"]:
            status_icon = "⏳"
        elif status in ["error", "failed"]:
            status_icon = "❌"

        # Format display name - use class name if available, otherwise use original file name
        display_name = class_name if class_name else original_file_name

        # Format uploader email - show only the part before @ symbol
        uploader_display = (
            uploader_email.split("@")[0] if "@" in uploader_email else uploader_email
        )

        # Format date to be more concise
        date_display = (
            upload_time.strftime("%Y-%m-%d")
            if isinstance(upload_time, datetime)
            else str(upload_time)
        )

        with st.expander(
            f"{status_icon} {display_name} | by {uploader_display} | {date_display}",
            expanded=False,
        ):
            description = item.get("description", "")
            size = item.get("formatted_size", "")

            if class_name:
                st.write(f"**Class**: {class_name}")
                st.write(f"**File**: {original_file_name}")
            else:
                st.write(f"**File**: {original_file_name}")

            if description:
                st.write(f"**Description**: {description}")
            st.write(f"**Uploaded**: {upload_time_str}")
            st.write(f"**Size**: {size}")

            # Audio player
            audio_url_with_sas = get_sas_url_for_audio_file_name(row_key)
            if audio_url_with_sas:
                st.audio(audio_url_with_sas)

            # Only fetch full transcript details if status is completed and user expands the item
            transcript = None
            if status == "completed" and transcript_id:
                try:
                    transcript = aai.Transcript.get_by_id(transcript_id)
                except pydantic.ValidationError as ve:
                    st.warning(
                        "Some transcript features may be limited due to API changes. Basic transcript text is still available."
                    )
                    logging.warning(
                        f"Validation error for transcript {transcript_id}: {str(ve)}"
                    )
                    try:
                        transcript = aai.Transcript.get_by_id(transcript_id)
                        if not hasattr(transcript, "text") or not transcript.text:
                            transcript = None
                    except Exception as e:
                        logging.error(f"Error getting transcript data: {str(e)}")
                        transcript = None
                except Exception as e:
                    st.error(f"Error loading transcript: {str(e)}")
                    logging.error(
                        f"Error loading transcript {transcript_id}: {str(e)}",
                        exc_info=True,
                    )

                if DEBUG and transcript:
                    # Good transcriptions have text and utterances
                    if transcript.text and transcript.utterances:
                        # Add download buttons in a row
                        col1, col2 = st.columns([1, 1])

                        with col1:
                            # Create download button for markdown
                            full_markdown = generate_transcript_markdown(transcript)
                            st.download_button(
                                label="Download as Markdown",
                                data=full_markdown,
                                file_name=f"{original_file_name}.md",
                                mime="text/markdown",
                                key=f"download_transcript_md_{row_key}",
                            )

                        with col2:
                            # Create download button for docx
                            docx_bytes = generate_transcript_docx(transcript)
                            st.download_button(
                                label="Download as Word",
                                data=docx_bytes,
                                file_name=f"{original_file_name}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_transcript_docx_{row_key}",
                            )

                        ### Show transcript
                        st.markdown("#### 📝 Transcript")
                        full_markdown = generate_transcript_markdown(transcript)
                        st.markdown(full_markdown)

                    elif transcript.text:
                        st.info("AI failed to distinguish speakers.")

                        # Add download buttons in a row
                        col1, col2 = st.columns([1, 1])

                        with col1:
                            # Create download button for text
                            st.download_button(
                                label="Download as Text",
                                data=transcript.text,
                                file_name=f"{original_file_name}.txt",
                                mime="text/plain",
                                key=f"download_transcript_txt_{row_key}",
                            )

                        with col2:
                            # Create download button for docx
                            docx_bytes = generate_transcript_docx(transcript)
                            st.download_button(
                                label="Download as Word",
                                data=docx_bytes,
                                file_name=f"{original_file_name}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_transcript_docx_no_speakers_{row_key}",
                            )

                        # Show full transcript
                        st.markdown("#### 📝 Transcript")
                        st.write(transcript.text)

            elif status in ["queued", "processing"]:
                with st.container(border=True):
                    st.markdown(
                        """
                        #### ⏳ Processing
                        The transcript is still being generated. This typically takes 1-2 minutes.
                        """
                    )
                    if st.button(
                        "Refresh Status",
                        icon="🔄",
                        key=f"refresh_status_body_{row_key}",
                    ):
                        st.rerun()

            elif status in ["error", "failed"]:
                with st.container(border=True):
                    st.markdown(
                        """
                        #### ❌ Error
                        There was a problem processing this transcript. Please try uploading the file again.
                        """
                    )

    except pydantic.ValidationError:
        # Handle case where transcript data is invalid/missing
        st.error(
            f"Could not load transcript {transcript_id} - the data may be corrupted or deleted"
        )
        logging.error(f"Failed to parse transcript data for ID: {transcript_id}")

    except Exception as e:
        # Handle any other errors including AssemblyAI API errors
        st.error(f"Error loading transcript: {str(e)}")
        logging.error(
            f"Error loading transcript {transcript_id}: {str(e)}", exc_info=True
        )


def display_table_data():
    """Display the table data with progress indicators"""
    items_list = load_table_data(table_client)

    if not items_list:
        st.info("No files found in the system")
        return

    # Sort by timestamp (newest first)
    items_list.sort(key=lambda x: x.get("_timestamp", datetime.min), reverse=True)

    # Calculate pagination
    total_items = len(items_list)
    start_idx = 0
    end_idx = st.session_state.items_per_page

    # Display items in fragments
    for item in items_list[start_idx:end_idx]:
        with st.container():
            display_transcript_item(item)

    # Load More button
    if end_idx < total_items:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button(
                f"Load More ({total_items - end_idx} remaining)",
                use_container_width=True,
            ):
                st.session_state.items_per_page += 5
                st.rerun()

    # Show total count
    st.caption(f"Showing {min(end_idx, total_items)} of {total_items} transcripts")


display_table_data()
